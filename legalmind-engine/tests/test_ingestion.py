import os
import pytest
import asyncio
from app.core.stores import CaseContext
from app.modules.dominion import Dominion
from app.models import RunStatus

# Fixture to create test data
@pytest.fixture(scope="session")
def test_data(tmp_path_factory):
    data_dir = tmp_path_factory.mktemp("data")

    # Create dummy DOCX
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is a test paragraph for LegalMind.")
    doc.add_paragraph("Another paragraph with some claims.")
    docx_path = data_dir / "test.docx"
    doc.save(docx_path)

    # Create dummy PDF
    from reportlab.pdfgen import canvas
    pdf_path = data_dir / "test.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "This is a test PDF for LegalMind.")
    c.drawString(100, 700, "Second line of text.")
    c.save()

    return str(docx_path), str(pdf_path)

@pytest.fixture
def case_context(tmp_path):
    case_path = tmp_path / "test_case_ingestion"
    os.makedirs(case_path, exist_ok=True)
    return CaseContext("test_case_ingestion", base_storage_path=str(tmp_path))

@pytest.fixture
def dominion(case_context):
    return Dominion(case_context)

@pytest.mark.asyncio
async def test_ingest_docx(dominion, test_data):
    docx_path, _ = test_data

    # Start job
    initial_state = await dominion.workflow_ingest_case(docx_path)
    assert initial_state.status == RunStatus.RUNNING
    run_id = initial_state.run_id

    # Poll for completion
    final_state = None
    for _ in range(30):
        await asyncio.sleep(0.5)
        state = dominion.get_job_status(run_id)
        if state and state.status in [RunStatus.COMPLETE, RunStatus.FAILED]:
            final_state = state
            break

    assert final_state is not None
    assert final_state.status == RunStatus.COMPLETE, f"Job failed with warnings: {final_state.warnings}"
    assert final_state.items_processed > 0

    # Verify Vault: File copied?
    vault_files = os.listdir(dominion.case_context.vault.vault_path)
    assert len(vault_files) >= 1

    # Verify Ledger
    segments = dominion.case_context.ledger.get_all_segments()
    assert len(segments) > 0
    assert "LegalMind" in segments[0].text

    # Verify Index (Chunks)
    chunks = dominion.case_context.index.get_all_chunks()
    assert len(chunks) > 0
    assert chunks[0].source == segments[0].source_asset_id

    # Verify Preservation (Chroma)
    assert dominion.preservation.collection.count() == len(chunks)

@pytest.mark.asyncio
async def test_ingest_pdf(dominion, test_data):
    _, pdf_path = test_data

    # Start job
    initial_state = await dominion.workflow_ingest_case(pdf_path)
    assert initial_state.status == RunStatus.RUNNING
    run_id = initial_state.run_id

    # Poll for completion
    final_state = None
    for _ in range(30):
        await asyncio.sleep(0.5)
        state = dominion.get_job_status(run_id)
        if state and state.status in [RunStatus.COMPLETE, RunStatus.FAILED]:
            final_state = state
            break

    assert final_state is not None
    assert final_state.status == RunStatus.COMPLETE, f"Job failed with warnings: {final_state.warnings}"
    assert final_state.items_processed > 0

    segments = dominion.case_context.ledger.get_all_segments()
    assert len(segments) > 0
    # Search for text in any segment (might be split differently)
    found_text = any("LegalMind" in s.text for s in segments)
    assert found_text

    chunks = dominion.case_context.index.get_all_chunks()
    assert len(chunks) > 0

    assert dominion.preservation.collection.count() == len(chunks)

@pytest.mark.asyncio
async def test_audit_workflow(dominion, test_data):
    docx_path, _ = test_data

    # 1. Ingest first to populate index
    ingest_state = await dominion.workflow_ingest_case(docx_path)
    # Wait for ingest
    for _ in range(30):
        await asyncio.sleep(0.5)
        state = dominion.get_job_status(ingest_state.run_id)
        if state and state.status == RunStatus.COMPLETE:
            break

    # 2. Run Audit
    audit_state = await dominion.workflow_audit_brief(docx_path)
    assert audit_state.status == RunStatus.RUNNING

    # Poll
    final_state = None
    for _ in range(30):
        await asyncio.sleep(0.5)
        state = dominion.get_job_status(audit_state.run_id)
        if state and state.status in [RunStatus.COMPLETE, RunStatus.FAILED]:
            final_state = state
            break

    assert final_state is not None
    assert final_state.status == RunStatus.COMPLETE, f"Audit failed: {final_state.warnings}"
    assert "report_path" in final_state.result_payload
    assert os.path.exists(final_state.result_payload["report_path"])
