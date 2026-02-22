import pytest
import os
import asyncio
from unittest.mock import MagicMock, patch
from app.core.stores import CaseContext
from app.modules.dominion import Dominion
from app.models import RunStatus

"""
Tests for DOCX ingestion granularity.
This test suite mocks heavy dependencies (Preservation/SentenceTransformers)
to focus on the logic of converting DOCX paragraphs into EvidenceSegments.
"""

# Fixture to create test data
@pytest.fixture(scope="session")
def test_docx(tmp_path_factory):
    data_dir = tmp_path_factory.mktemp("data_repro")
    from docx import Document
    doc = Document()
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("Paragraph 2")
    doc.add_paragraph("Paragraph 3")
    docx_path = data_dir / "repro.docx"
    doc.save(docx_path)
    return str(docx_path)

@pytest.fixture
def dominion(tmp_path):
    case_path = tmp_path / "test_case_repro"
    os.makedirs(case_path, exist_ok=True)
    ctx = CaseContext("test_case_repro", base_storage_path=str(tmp_path))

    # Patch Preservation to avoid heavy dependencies like sentence_transformers
    with patch('app.modules.dominion.Preservation') as MockPreservation:
        instance = MockPreservation.return_value
        instance.dense_indexer = MagicMock()
        instance.bm25_indexer = MagicMock()

        dom = Dominion(ctx)
        return dom

@pytest.mark.asyncio
async def test_ingest_docx_chunking(dominion, test_docx):
    """
    Verifies that DOCX ingestion creates a separate segment/chunk for each paragraph.
    Previous behavior was creating a single segment for the entire document.
    """
    # Start job
    initial_state = await dominion.workflow_ingest_case(test_docx)
    run_id = initial_state.run_id

    # Poll for completion
    final_state = None
    for _ in range(30):
        await asyncio.sleep(0.5)
        state = dominion.get_job_status(run_id)
        if state and state.status in [RunStatus.COMPLETE, RunStatus.FAILED]:
            final_state = state
            break

    assert final_state is not None
    assert final_state.status == RunStatus.COMPLETE, f"Job failed: {final_state.warnings}"

    # Check chunks processed
    print(f"Items Processed: {final_state.items_processed}")

    # Retrieve segments directly to verify count
    segments = dominion.case_context.ledger.get_all_segments()
    print(f"Segments Count: {len(segments)}")
    for s in segments:
        print(f"Segment: {s.text[:50]}...")

    # Desired expectation: 3 chunks, 3 segments
    assert final_state.items_processed == 3
    assert len(segments) == 3

    # Verify locations
    locations = sorted([s.location for s in segments])
    assert locations == ["para_1", "para_2", "para_3"]
