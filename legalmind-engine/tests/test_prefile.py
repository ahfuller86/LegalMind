import os
import pytest
import asyncio
from docx import Document
from app.core.stores import CaseContext
from app.modules.dominion import Dominion
from app.models import RunStatus, FilingRecommendation

@pytest.fixture
def dominion(tmp_path):
    case_context = CaseContext("test_case_prefile", base_storage_path=str(tmp_path))
    return Dominion(case_context)

@pytest.fixture
def brief_path(tmp_path):
    d = tmp_path / "data"
    d.mkdir()
    p = d / "brief.docx"
    doc = Document()
    doc.add_paragraph("Argument 1: The sky is blue.")
    doc.add_paragraph("See 347 U.S. 483.") # Brown v. Board (Verified)
    doc.add_paragraph("Argument 2: The moon is made of cheese.")
    doc.save(p)
    return str(p)

@pytest.mark.asyncio
async def test_cite_check_workflow(dominion, brief_path):
    # Read text from brief manually for this test or use helper
    doc = Document(brief_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    state = await dominion.workflow_cite_check(text)
    assert state.status == RunStatus.RUNNING

    # Poll
    final_state = None
    for _ in range(30):
        await asyncio.sleep(0.1)
        s = dominion.get_job_status(state.run_id)
        if s.status in [RunStatus.COMPLETE, RunStatus.FAILED]:
            final_state = s
            break

    assert final_state.status == RunStatus.COMPLETE
    assert final_state.result_payload["findings_count"] == 1
    assert os.path.exists(final_state.result_payload["report_path"])

@pytest.mark.asyncio
async def test_prefile_gate_workflow(dominion, brief_path):
    # Ingest evidence first to have something to retrieve
    # (Using mock/stub ingestion or creating index directly for speed)
    # But for now, we rely on empty index -> 0 supported claims is fine

    state = await dominion.workflow_prefile_gate(brief_path)
    assert state.status == RunStatus.RUNNING

    final_state = None
    for _ in range(50):
        await asyncio.sleep(0.1)
        s = dominion.get_job_status(state.run_id)
        if s.status in [RunStatus.COMPLETE, RunStatus.FAILED]:
            final_state = s
            break

    assert final_state.status == RunStatus.COMPLETE
    result = final_state.result_payload
    gate = result["gate_result"]

    # We expect CLEAR or REVIEW based on mock logic
    # Brown v Board is Verified -> Good
    # Claims -> No index -> Not Supported -> Risk Score increases
    # Sentinel logic: Not Supported adds 5.0 risk.
    # 2 claims?

    assert gate["filing_recommendation"] in ["CLEAR", "REVIEW_REQUIRED"]
    assert gate["citation_summary"]["verified"] == 1
    assert os.path.exists(result["report_path"])
